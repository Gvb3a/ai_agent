import os
import requests
import json
import re
from urllib.parse import quote
from datetime import datetime
from bs4 import BeautifulSoup
from typing import Literal
import aiohttp
import asyncio
import base64
import hashlib
import PIL.Image


from ..config.logger import logger
from ..config.config import load_config


config = load_config()


def date_hash() -> str:
    return hashlib.md5(datetime.now().strftime("%Y-%m-%d %H:%M:%S.%f").encode()).hexdigest()


def download_image(url: str) -> str:
    'Downloads the image from the link. Returns the name of the downloaded image (name is time).'
    response = requests.get(url)
    file_name = f'{hash(url)}.png'
    with open(file_name, 'wb') as file:
        file.write(response.content)
    logger.info(file_name)
    return file_name



async def async_download_image(session, url, name):
    save_path = name + '.png'
    try:
        async with session.get(url) as response:
            if response.status == 200:
                with open(save_path, 'wb') as file:
                    file.write(await response.read())
                return save_path
            else:
                logger.error(f"Failed to download {url}: HTTP {response.status}")
    except Exception as e:
        logger.error(f"Error downloading {url}: {e}", exc_info=True)


async def download_images(image_urls):
    name = date_hash()
    async with aiohttp.ClientSession() as session:
        tasks = []
        for i, url in enumerate(image_urls):
            try:
                tasks.append(async_download_image(session, url, f'{name} {i}'))
            except:
                pass
        
        return await asyncio.gather(*tasks)
    

def merge_pngs_vertically(image_paths: list) -> str:
    output_path = date_hash() + '.png'
    images = [PIL.Image.open(img_path) for img_path in image_paths]

    total_width = max(img.width for img in images)

    total_height = sum(img.height for img in images)

    new_image = PIL.Image.new("RGBA", (total_width, total_height), (255, 255, 255, 255))

    current_height = 0

    for img in images:
        x_offset = (total_width - img.width) // 2
        new_image.paste(img, (x_offset, current_height))
        current_height += img.height

    new_image.save(output_path)
    return output_path
    



# =========================< FILE TO TEXT >=========================
# pdf, docx, mp3 to text. TODO: summary in large document (or embbeding)
import PyPDF2
import docx
from groq import Groq

groq_api_key = config.api.groq_key
groq_client = Groq(api_key=groq_api_key)


def speech_recognition(file_name: str) -> str:  # TODO: local_whisper

    with open(file_name, "rb") as file:
        translation = groq_client.audio.transcriptions.create(
        file=(file_name, file.read()),
        model="whisper-large-v3")
        
        text = translation.text
    
    return str(text).strip()


def pdf_to_text(pdf_path: str) -> str:
    with open(pdf_path, 'rb') as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        extracted_text = ""
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                extracted_text += text

    return extracted_text


def docx_to_text(docx_path: str) -> str:
    doc = docx.Document(docx_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)

    return '\n'.join(full_text)


def files_to_text(files: list | str) -> str:
    'Turns text files or audio into text'
    if type(files) == str:
        files = [files]

    result = ''
    
    try:
        for file_path in files:

            if file_path.endswith('.pdf'):
                text = pdf_to_text(file_path)
            
            elif file_path.endswith('.docx'):
                text = docx_to_text(file_path)
            
            elif file_path.endswith('.mp3'):
                text = speech_recognition(file_path)

            else:
                with open(file_path, 'r') as file:
                    text = file.read()
            
            result += f'{file_path}:\n{text}\n\n'
            
        return result
    
    except Exception as e:
        logger.error(e, exc_info=True)
        return f'Error reading {file_path}'



# =========================< CONVERT FILES >=========================
import fitz
def pdf_to_image(pdf_path: str, quality: int = 3, extension: str = 'png') -> list[str]:
    file_name = pdf_path[:-4]
    doc = fitz.open(pdf_path)
    photos = []
    count = len(doc)
        
    for i in range(count):
        page = doc.load_page(i)
        pix = page.get_pixmap(matrix=fitz.Matrix(quality, quality))
        temp_file_name = f"{file_name}_{i}.{extension}"
        pix.save(temp_file_name)
        photos.append(temp_file_name)

    doc.close()
    logger.info(f'Success({count} pages): {photos}')
    return photos


def merges_pdf(files: list[str]) -> str:
    merger = fitz.Document()

    for file in files:
        doc = fitz.open(file)
        merger.insert_pdf(doc)
        doc.close()

    merged_file = f'{files[0][:-4]}_merged.pdf'
    merger.save(merged_file)

    logger.info(f'Success({len(files)} files): {merged_file}')
    return merged_file




# =========================< GROQ API >=========================
def groq_api(messages: list, model: str = 'llama-3.3-70b-versatile') -> str:
    # https://console.groq.com/docs/models
    response = groq_client.chat.completions.create(
        messages=messages,
        model=model)
        
    return str(response.choices[0].message.content)



# =========================< GOOGLE API >=========================
import google.generativeai as genai

os.environ["GOOGLE_API_KEY"] = config.api.gemini_key
genai.configure(api_key=os.environ['GOOGLE_API_KEY'])
google_model_thinking = genai.GenerativeModel("gemini-2.0-flash-thinking-exp")
google_model = genai.GenerativeModel("gemini-2.5-flash")

def genai_api(messages: list[dict] | str, file_paths: list | str = []) -> str:
    # TODO:  model: Literal['google', 'google_flash', 'google_thinking'] = 'google_flash'
    if type(messages) == str:
        user_message = messages
        formatted_history = []

    else:  # google is a bit special and to keep everything in the same style I do the transformation here.
        user_message = messages[-1]['content']
        history = messages[:-1]

        formatted_history = []
        for message in history:
            role = 'user' if message['role'] in ['user', 'system'] else 'model' 
            content = message['content']
            formatted_history.append({"role": role, "parts": content})


    if type(file_paths) == str:
        file_paths = [file_paths]


    chat = google_model.start_chat(history=formatted_history)

    files = []
    for file_path in file_paths:
        if file_path.endswith('.png') or file_path.endswith('.jpg') or file_path.endswith('.jpeg') or file_path.endswith('.webp'):
            files.append(PIL.Image.open(file_path))

        else:
            user_message += '\n\n' + file_path + ':\n' + files_to_text(file_path)

    if file_paths:
        response = chat.send_message([user_message] + files)

    else:
        response = chat.send_message(user_message)


    return response.text


# =========================< LLM API >=========================
def llm_api(messages: list[dict], files: str | list = [], provider: Literal['groq', 'google'] = 'google'):
    # provider replace to model
    if type(messages) == str:
        messages = [{'role': 'user', 'content': messages}]

    if type(files) == str:
        files = [files]
        
    start_time = datetime.now()
    
    if provider == 'google' or files:
        try:
            answer = genai_api(messages, files)
        except Exception as e:
            logger.error(f'google error: {e}', exc_info=True)
            answer = groq_api(messages)
    
    else:
        answer = groq_api(messages)

    logger.info(f'answer: {answer}, messages: {messages}, files: {files}, provider: {provider}, time: {datetime.now()-start_time}')

    return answer



# =========================< WOLFRAM ALPHA >=========================
WOLFRAM_SIMPLE_API = config.api.wolfram_simple_key
WOLFRAM_SHOW_STEPS_RESULT = config.api.wolfram_full_key


def calculator(expression: str) -> str:
    try:
        expression = expression.replace('^', '**')
        logger.info(expression)
        return str(eval(expression))
    
    except Exception as e:
        logger.error(f"Calculator error: {e}", exc_info=True)
        return f"Calculator error: {e}"


def wolfram_short_answer_api(text: str) -> str:
    'https://products.wolframalpha.com/short-answers-api/documentation'
    query = quote(text)
    url = f'https://api.wolframalpha.com/v1/result?appid={WOLFRAM_SIMPLE_API}&i={query}'
    answer = requests.get(url).text
    logger.info(answer)
    return answer


def wolfram_llm_api(text: str) -> tuple[str, list]:
    'https://products.wolframalpha.com/llm-api/documentation - text version of the WolframAlpha page answer. Returns the text and a list of links to images'
    query = quote(text)
    url = f'https://www.wolframalpha.com/api/v1/llm-api?input={query}&appid={WOLFRAM_SHOW_STEPS_RESULT}'
    answer = requests.get(url).text
    # TODO: improve (add async) and add llm
    answer = answer[:answer.find('Wolfram|Alpha website result for "')]
    links = re.findall(r'https?://\S+', answer)
    answer = re.sub(r'https?://\S+', 'Images will be attached to the answer', answer)
    logger.info(f'{answer}, {links}')
    return answer, links


def wolfram_simple_api(text: str) -> str:
    'https://products.wolframalpha.com/simple-api/documentation - WolframAlpha answer page image'
    query = quote(text)
    link = f'https://api.wolframalpha.com/v1/simple?appid={WOLFRAM_SIMPLE_API}&i={query}%3F'
    file_name = download_image(link)
    logger.info(file_name)
    return file_name


def wolfram_short_answer(query: str) -> str:
    'Short answer from WolframAlpha'
    quick_answer = wolfram_short_answer_api(query)
    logger.info(quick_answer)
    return quick_answer


def wolfram_full_answer(text: str):  # TODO: async + wolfram_simple_api
    'returns the text version of the answer sheet, the image links and the answer sheet as a picture'

    full_answer, full_answer_images = wolfram_llm_api(text)
        
    images = full_answer_images # TODO: [wolfram_simple_api(text)] + full_answer_images 
    logger.info(f'{full_answer}, {images}')
    return full_answer, images



# =========================< TAVILY AND DUCKDUCKGO (INTERNERT) >=========================
from tavily import TavilyClient
from duckduckgo_search import DDGS
from deep_translator import GoogleTranslator, single_detection

detect_language_api_key = config.api.detect_language_key
tavily_client = TavilyClient(api_key=config.api.tavily_key)


def detect_language(text: str) -> str:  # TODO: ggcs translator
    try:
        return single_detection(text=text, api_key=detect_language_api_key, detailed=False)
    except:
        return 'en'
    

def translate(text: str, target_language: str = 'en', source_language: str = 'auto') -> str:
    translated_text = GoogleTranslator(source=source_language, target=target_language).translate(text)
    return translated_text


def parsing(links: str | list) -> str:
    start_time = datetime.now() 
    try:
        responce = tavily_client.extract(urls=links)
        return [r['raw_content'] for r in responce['results']]
    
    except Exception as e:
        if type(links) == str:
            links = [links]
        result = ''
        for link in links:
            try:
                responce = tavily_client.extract(urls=link)
                result += f'{link}: {responce["result"][0]["raw_content"]}\n'
            except:
                result += f'{link}: Error when extracting text ({e})\n'

    logger.info(f'{links}, {result}, {datetime.now() - start_time}')
    return result
    

def DDGS_answer(text: str) -> str:
    'A short answer like Google. Sometimes nothing comes up. But mostly the answer comes from wikipedia.'
    try:
        response = DDGS().answers(text)[:3]
        return '\n'.join([r["text"] for r in response])
    except:
        return ''


def DDGS_images(text: str, max_results: int = 9) -> list[str]:
    images = [i['image'] for i in DDGS().images(text, max_results=max_results)]
    return images


def tavily_search(query: str, max_results: int = 7):
    response = tavily_client.search(query=query, max_results=max_results)
    result = ''
    for i in response['results']:
        result += f'{i["url"]}({i["title"]}): {i["content"]}\n'
    return result


def tavily_get_context(query: str, topic = 'news') -> str:
    return tavily_client.get_search_context(query=query, topic=topic).replace('\\', '')


def tavily_content(text: str, max_results: int = 4):
    response = tavily_client.search(query=text, search_depth='basic', max_results=max_results)['results']
    results = ''
    for r in response:
        results += f'{r["url"]}: {r["content"]}\n'

    return results


prompt_for_sum = f"""You are a precise summarization expert. Your task is to create a clear and relevant summary based on the provided text and query.

Context: I will provide you with:
1. Query: A specific question or topic of interest
2. Source text: Content from a webpage or document or a collection of several responses

Ignore unnecessary information and answer the query well.
"""


def sum_page(link: str, query: str) -> str:
    prompt = prompt_for_sum + f'\n\nQuery:\n{query}\n\nSource text:\n{parsing(link)}'  # TODO: IMPROVE
    messages = [{"role": "user", "content": prompt}]
    try:
        llm_answer = llm_api(messages=messages, provider='google')
        logger.info(llm_answer)
    except Exception as e:
        l = len(str(messages))
        llm_answer = 'Error'
        logger.error(f'{e}. len: {l}', exc_info=True)

    return f'{link}: {llm_answer}\n\n\n'


def google_links(text: str, max_results: int = 5) -> list[str]:
    try:
        links = [i['href'] for i in DDGS().text(text, max_results=max_results)]
        logger.info(str(links))
        return links
    except Exception as e:
        logger.error(f'{e}', exc_info=True)
        return []


def google_short_answer(text: str) -> str:
    resp = DDGS_answer(text)
    final_answer = resp if resp else tavily_search(text)
    logger.info(final_answer)
    return final_answer
    

def google_full_answer(text: str, max_results: int = 5):  # TODO: async, for lyrics

    start_time = datetime.now()
    links = google_links(text=text, max_results=max_results)

    sum_answers = ''
    for link in links:
        sum_answers += sum_page(link=link, query=text)  # TODO: max_result excluding parsing error

    prompt = prompt_for_sum + f"""Summarize the text above in a concise and relevant way. Ensure that the summary is well-organized and captures the main points of the text. The summary should be based on the query and the provided text. If the text is not relevant to the query, please provide a summary that is not relevant to the query. Source text will be composed of several responses. Write the final text"""
    prompt +=  f'\n\nQuery:\n{text}\n\nSource text:\n{sum_answers}'
    
    final_messages = [{"role": "user", "content": prompt}]
    final_answer = llm_api(messages=final_messages, provider='groq')

    logger.info(f'{final_answer}, {datetime.now() - start_time}')
    return final_answer


def google_news(text: str):
    news_content = tavily_get_context(text)
    
    prompt = prompt_for_sum + f"""Turn all these texts into one\n\nQuery:\n{text}\n\nSource text:\n{news_content}"""

    final_messages = [{"role": "user", "content": prompt}]
    final_answer = llm_api(messages=final_messages, provider='groq')

    logger.info(final_answer)
    return final_answer



async def google_image(text, max_results=9, download_images_or_not=True):
    urls = DDGS_images(text, max_results=max_results)
    start_time = datetime.now()
    

    if not download_images_or_not:
        logger.info(urls)
        return f'google_image: The {len(urls)} of images on the {text} query will be prefixed to your response', urls

    
    file_paths = await download_images(urls)
    
    logger.info(f'{file_paths}, {datetime.now()-start_time}')
    return f'The {len(file_paths)} of the {text} query images will be appended to the response. Answer other user questions, tell information about the query, or say something like images found. DON\'T write anything like [Image of ...] or you\'ll be shut down.', file_paths


# =========================< YOUTUBE >=========================
from youtube_transcript_api import YouTubeTranscriptApi

def transcript2text(transcript: list[dict], sep='\n'):
    result = []
    current_minute = -1

    for t in transcript:
        minute = int(t["start"] // 60)
        if minute != current_minute:
            result.append(f"{minute} minute:")
            current_minute = minute
        
        result.append(t["text"])

    return sep.join(result)



def get_youtube_transcripts(link: str, language: str = 'en'):

    if 'youtube.com' in link:
        # https://www.youtube.com/watch?v=xxxxxxxxxxx -> xxxxxxxxxxx
        pattern_for_youtube_video = r"(?:v=|\/)([a-zA-Z0-9_-]{11})"  # unreadable, but it works perfectly
        match = re.search(pattern_for_youtube_video, link)
        if match:
            video_id = match.group(1)
        else:
            return 'Error: incorrect link'
    else:
        video_id = link

    
    # get title
    response = requests.get(link)
    soup = BeautifulSoup(response.text, 'html.parser')
    title_tag = soup.find("title")
    if title_tag:
        title = title_tag.text.replace(" - YouTube", "")
    else:
        title = 'Error: title not found'

    # get transcript: [{'text': str, 'start': float, 'duration': float}, ...]
    try:  
        transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=[language])
        logger.info(f'1 (defult language): {language}, {link}({title})')
        return transcript2text(transcript), title
    except:
        pass


    try:
        language = detect_language(title)
        transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=[language])
        logger.info(f'2 (detect language): {language}, {link}({title})')
        return transcript2text(transcript), title
    except:
        pass


    try:
        transcripts = YouTubeTranscriptApi.list_transcripts(video_id)
        for transcript in transcripts:
            transcript = transcript.translate('en').fetch()
            break
        logger.info(f'3 (translate). {link}({title})')
    except Exception as e:
        logger.error(f'error: {e}', exc_info=True)
        transcript = [{'text': f'Subtitles are not available', 'start': 0.0}]
    
    return transcript2text(transcript), title


def youtube_sum(link: str, question: str | None = None, language: str = 'en') -> str: 
    # make it possible to ask questions.
    text, title = get_youtube_transcripts(link, language)

    if question:
        content = f'Answer the question "{question}" based on this YouTube video "{title}": {text}'
    else:
        content = f'Summarize this YouTube video "{title}": {text}'

    try:
        answer = llm_api(messages=[{'role': 'user', 'content': content}], provider='google')
        logger.info(f'{link} ({question}): {answer}')
        return answer
    except Exception as e:
        logger.error(f'{link} error ({e})', exc_info=True)
        return f'Error({e}). Most like too much text or a completely incomprehensible error that cannot be fixed (or subtitles are not available). Tell the user to try again later'



# =========================< CODE INTERPRETER (eb2) >=========================
from e2b_code_interpreter import Sandbox

def code_interpreter(code: str):
    with Sandbox() as sandbox:
        execution = sandbox.run_code(code)

        stdout = execution.logs.stdout
        stderr = execution.logs.stderr

        try:
            first_result = execution.results[0]
            if first_result.png:
                image_file_name = 'e2b_image.png'
                with open(image_file_name, 'wb') as f:
                    f.write(base64.b64decode(first_result.png))
            else:
                image_file_name = None
        except:
            image_file_name = None

        try:
            result = stdout[0]
        except:
            result = ''

        try:
            result += '\n' + stderr[0]
        except:
            pass
        
        logger.info(f'{result}, {image_file_name}')
        return result, image_file_name


# =========================< LATEX >=========================
def latex_expression_to_png(expression: str, size: int = 400):
    try:
        expression = expression.strip('$')
        link = 'https://latex.codecogs.com/png.image?\\dpi{' + str(size) + '}' + expression.replace(' ', '%20')
        response = requests.get(link)
        file_name = hashlib.md5(expression.encode()).hexdigest() + '.png'
        if response.status_code == 200:
            with open(file_name, 'wb') as f:
                f.write(response.content)
            logger.info(expression)
            return f'A rendering of the expression {expression} will be added to the answer', [file_name]
        else:
            logger.error(expression)
            return None
    except:
        logger.error(expression, exc_info=True)
        return None
    




async def async_expressions_to_png(expressions: list[str], size: int = 400):
    try:
        links = []
        for expression in expressions:
            if type(expression) != str:
                l = [i for i in expression if i]
                expression = l[0]
            expression = expression.strip('$')
            if len(expression) <= 3:
                continue
            link = 'https://latex.codecogs.com/png.image?\\dpi{' + str(size) + '}' + expression.replace(' ', '%20')
            links.append(link)
        logger.info(f'expressions: {expressions}')
        return await download_images(links)
    except Exception as e:
        logger.error(f'error: {e}', exc_info=True)
        return []
    

def latex_to_pdf(content: str, recursion_turn: int = 1) -> str | None:


    if '\\begin{document}' not in content:
        preambula = '''\\documentclass[a4paper]{article}
\\usepackage[english,russian]{babel}
\\usepackage[utf8]{inputenc}
\\usepackage{geometry}
\\geometry{a4paper, left=15mm, right=15mm, top=15mm, bottom=17mm}
\\usepackage{amsmath, amssymb, amsfonts}
'''
        content = preambula + '\\begin{document}\n' + content + '\n\\end{document}'
    if recursion_turn > 3:
        return None
    
    link = 'https://latexonline.cc/compile?text=' + quote(content)
    response = requests.get(link)
    file_name = hashlib.md5(content.encode()).hexdigest() + '.pdf'
    if response.status_code == 200:
        with open(file_name, 'wb') as f:
            f.write(response.content)
        logger.info(content)
        return file_name
    else:
        logger.error(f'Error: {response.text}, recursion_turn: {recursion_turn}, content: {content}')
        prompt = 'You are a LaTeX expert. You have to fix the LaTeX Document so that the error does not occur (if the error is very unclear, you can remove the part of the text with the error). '\
                 'Your whole reply will go in the reply, so you can write your thoughts in the comments (after %), nobody will see them. '\
                f'Error: {response.text}\n\nText:{content}'
        answer = llm_api(prompt, provider='google')
        if answer.startswith('```latex'):
            answer = answer[8:-4]
        elif answer.startswith('```'):
            answer = answer[3:-4]
        logger.info(f'Ask llm: {answer}. ')
        return latex_to_pdf(answer, recursion_turn+1)
    

def text_to_pdf_document(text: str):
    
    prompt = "You are a LaTeX expert. Your task is to convert this text in MarkDown markup into a LaTeX document. You don't have to cause an error.  "\
             "Tips: use in preambula \\documentclass[a4paper]{article}, \\usepackage[english,russian]{babel}, \\usepackage[utf8]{inputenc}, \\usepackage{geometry}, \\geometry{a4paper, left=15mm, right=15mm, top=15mm, bottom=17mm}, \\usepackage{amsmath, amssymb}. You can add your own."\
             "You encapsulate EACH math equation IN $$ and use LaTeX. If something like V2 or T₃ occurs in the text, convert it to V_2 and T_3, Δ to \\Delta and etc. For bold (**) use \\textbf{}, for bullet list use itemize and so on. "\
            f"Text:\n{text}"
    
    
    answer = llm_api(prompt, provider='google').strip()
    
    if answer.startswith('```latex'):
        answer = answer[8:-4]
    elif answer.startswith('```'):
        answer = answer[3:-4]

    file_name = latex_to_pdf(answer)
    logger.info(f'text: {text}, answer: {answer}', error=not bool(file_name))
    return file_name



# =========================< IMDB >=========================
from PyMovieDb import IMDB
imdb = IMDB()


def imdb_search(title):
    responce_json = imdb.search(title)
    responce_dict = json.loads(responce_json)
    result = responce_dict['results']
    return result



def imdb_get_film_by_id(imdb_id: str, download_image_or_not=True) -> tuple[str, list[str]]:
    responce_json = imdb.get_by_id(imdb_id)
    responce_dict = json.loads(responce_json)

    if responce_dict.get('status', False) == 404:
        return 'Error', []
    
    result = dict(responce_dict)
    
    result['url'] = f'https://www.imdb.com/title/{imdb_id}'


    if download_image_or_not:
        poster = download_image(responce_dict['poster'])
    else:
        poster = responce_dict['poster']

    del result['poster']

    result['review'] = responce_dict['review']['reviewBody']

    result['rating'] = responce_dict['rating']['ratingValue']
    result['ratingCount'] = responce_dict['rating']['ratingCount']

    return result, [poster]


def imdb_api(title):
    search_result = imdb_search(title)
    if len(search_result) == 0:
        logger.error(f'No result for {title}')
        return 'Error'
    
    imdb_id = search_result[0]['id']
    result, poster = imdb_get_film_by_id(imdb_id)

    logger.info(f'Succes result for {title}')
    return result, poster

