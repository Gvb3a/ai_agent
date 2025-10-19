import docx
import fitz
import PyPDF2
import asyncio
import aiohttp
import hashlib
import requests
import PIL.Image
from groq import Groq
from datetime import datetime
from ...config.logger import logger
from ...config.config import load_config


config = load_config()
groq_api_keys = config.api.groq_key
groq_client = [Groq(api_key=key) for key in groq_api_keys]


def date_hash() -> str:
    return hashlib.md5(datetime.now().strftime("%Y-%m-%d %H:%M:%S.%f").encode()).hexdigest()


def speech_recognition(file_name: str) -> str:  # TODO: local_whisper
    with open(file_name, "rb") as file:
        for client in groq_client:
            try:
                translation = client.audio.transcriptions.create(
                    file=(file_name, file.read()),
                    model="whisper-large-v3")
                text = translation.text
                logger.info(f'Success: {text}')
                break
            except Exception as e:
                logger.error(f'Error with speech recognition on {client.api_key}: {e}', exc_info=True)
                text = f'Error with speech recognition: {e}'
                continue
    
    return str(text).strip()


# =========================< DOWNLOAD IMAGE >=========================
def download_image(url: str) -> str:
    'Downloads the image from the link. Returns the name of the downloaded image'
    response = requests.get(url)
    if response.status_code == 200:
        file_name = f'{hash(url)}.png'
        with open(file_name, 'wb') as file:
            file.write(response.content)
        logger.info(file_name)
        return file_name
    else:
        logger.error(f"Failed to download {url}: HTTP {response.status_code}")
        return None
    

async def async_download_image(session, url, name):
    save_path = name + '.png'
    try:
        async with session.get(url) as response:
            if response.status == 200:
                with open(save_path, 'wb') as file:
                    file.write(await response.read())
                return save_path
            else:
                logger.error(f"Failed to download {url}: HTTP {response.status}")
    except Exception as e:
        logger.error(f"Error downloading {url}: {e}", exc_info=True)


async def download_images(image_urls):
    name = date_hash()
    async with aiohttp.ClientSession() as session:
        tasks = []
        for i, url in enumerate(image_urls):
            try:
                tasks.append(async_download_image(session, url, f'{name} {i}'))
            except:
                pass
        
        return await asyncio.gather(*tasks)
    


# =========================< FILE TO TEXT >=========================
def pdf_to_text(pdf_path: str) -> str:
    with open(pdf_path, 'rb') as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        extracted_text = ""
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                extracted_text += text

    return extracted_text


def docx_to_text(docx_path: str) -> str:
    doc = docx.Document(docx_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)

    return '\n'.join(full_text)


def files_to_text(files: list | str) -> str:
    'Turns text files or audio into text'
    if type(files) == str:
        files = [files]

    result = ''
    
    try:
        for file_path in files:

            if file_path.endswith('.pdf'):
                text = pdf_to_text(file_path)
            
            elif file_path.endswith('.docx'):
                text = docx_to_text(file_path)
            
            elif file_path.endswith('.mp3'):
                text = speech_recognition(file_path)

            elif file:
                with open(file_path, 'r') as file:
                    text = file.read()
            
            result += f'{file_path}:\n{text}\n\n---\n\n'
            
        return result
    
    except Exception as e:
        logger.error(e, exc_info=True)
        return f'Error reading {file_path}'



# =========================< WORK WITH FILES >=========================
def pdf_to_image(pdf_path: str, quality: int = 3, extension: str = 'png') -> list[str]:
    file_name = pdf_path[:-4]
    doc = fitz.open(pdf_path)
    photos = []
    count = len(doc)
        
    for i in range(count):
        page = doc.load_page(i)
        pix = page.get_pixmap(matrix=fitz.Matrix(quality, quality))
        temp_file_name = f"{file_name}_{i}.{extension}"
        pix.save(temp_file_name)
        photos.append(temp_file_name)

    doc.close()
    logger.info(f'Success({count} pages): {photos}')
    return photos


def merges_pdf(files: list[str]) -> str:
    merger = fitz.Document()

    for file in files:
        doc = fitz.open(file)
        merger.insert_pdf(doc)
        doc.close()

    merged_file = f'{files[0][:-4]}_merged.pdf'
    merger.save(merged_file)

    logger.info(f'Success({len(files)} files): {merged_file}')
    return merged_file


def merge_pngs_vertically(image_paths: list) -> str:
    output_path = date_hash() + '.png'
    images = [PIL.Image.open(img_path) for img_path in image_paths]

    total_width = max(img.width for img in images)

    total_height = sum(img.height for img in images)

    new_image = PIL.Image.new("RGBA", (total_width, total_height), (255, 255, 255, 255))

    current_height = 0

    for img in images:
        x_offset = (total_width - img.width) // 2
        new_image.paste(img, (x_offset, current_height))
        current_height += img.height

    new_image.save(output_path)
    return output_path


